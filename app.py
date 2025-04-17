from quiz_ai import transcribe, generate_quiz
from flask import Flask, redirect, request, jsonify, send_file, render_template, url_for
from flask_cors import CORS
import os
import uuid
import shutil
from docx import Document
import json
from io import BytesIO

import quiz_ai
 
app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/')
def index():
    quizzes_list = []
    for dir_name in os.listdir(UPLOAD_FOLDER):
        dir_path = os.path.join(UPLOAD_FOLDER, dir_name)
        for filename in os.listdir(dir_path):
            if filename.endswith(".json"):
                quiz_filepath = os.path.join(dir_path, filename)
                with open(quiz_filepath, "r") as f:
                    quiz_data = json.load(f)
                break
        quizzes_list.append({
            "dir_name" : dir_name,
            "title" : quiz_data[0]["title"]
        })

    return render_template('home.html', quizzes_list=quizzes_list)

@app.route('/upload', methods=['GET', 'POST'])
def upload_file():

    if request.method == "POST":
        files = request.files
        dir_name = str(uuid.uuid4())
        upload_dir = os.path.join(UPLOAD_FOLDER, dir_name)
        os.makedirs(upload_dir, exist_ok=True)

        for key in files:
            file = files[key]
            if file:
                filename = file.filename
                filepath = os.path.join(upload_dir, filename)
                file.save(filepath)

                if filename.lower().endswith(('.mp4', '.mkv', '.avi', '.mov', '.webm')):
                    transcribe(dir_name, filename)

        return redirect(url_for("quiz", dir_name=dir_name))
    return render_template('upload.html')

@app.route('/quiz/<dir_name>', methods=['GET'])
def quiz(dir_name):
    quiz_filepath = os.path.join(UPLOAD_FOLDER, dir_name, "quiz.json")
    if not os.path.exists(quiz_filepath):
        filepaths = []
        dir_path = os.path.join(UPLOAD_FOLDER, dir_name)
        for file in os.listdir(dir_path):
            if file.endswith(('.docx', '.pdf', '.doc', '.txt')):
                filepaths.append(os.path.join(dir_path, file))
        generate_quiz(dir_path, filepaths)
    
    return render_template("quiz.html", dir_name=dir_name)

@app.route('/quiz_data/<dir_name>', methods=['GET'])
def quiz_data(dir_name):
    quiz_filepath = os.path.join(UPLOAD_FOLDER, dir_name, "quiz.json")
    if not os.path.exists(quiz_filepath):
        return jsonify({"error": "Quiz file not found"}), 404
    
    with open(quiz_filepath, "r") as f:
        quiz_data = json.load(f)
    
    return jsonify(quiz_data)

@app.route('/download_quiz/<dir_name>')
def download_quiz(dir_name):
    try:
        quiz_path = os.path.join(UPLOAD_FOLDER, dir_name, "quiz.json")
        if not os.path.exists(quiz_path):
            return jsonify({"error": "Quiz not found"}), 404

        with open(quiz_path, 'r') as file:
            quiz_data = json.load(file)

        doc = Document()
        doc.add_heading(quiz_data[0]['title'], level=1)

        for item in quiz_data[1:]:
            doc.add_paragraph(f"{item['question']}", style='List Number')

            if item.get("options"):
                for opt in item["options"]:
                    doc.add_paragraph(f"{opt}", style='List Bullet')

            answer_paragraph = doc.add_paragraph()
            answer_run = answer_paragraph.add_run(f"Answer: {item['answer']}")
            answer_run.bold = True

            doc.add_paragraph()

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        filename = quiz_data[0]['title'].replace(' ', '_') + '.docx'
        return send_file(
            doc_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)